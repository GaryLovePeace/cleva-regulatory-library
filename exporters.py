from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Iterable, Mapping, Any

from docx import Document
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


def _json_list(value: Any) -> str:
    if not value:
        return ""
    if isinstance(value, str):
        try:
            value = json.loads(value)
        except json.JSONDecodeError:
            return value
    return ", ".join(str(v) for v in value)


def regulations_to_xlsx(rows: Iterable[Mapping[str, Any]]) -> bytes:
    rows = list(rows)
    wb = Workbook()
    ws = wb.active
    ws.title = "Regulation_Master"
    headers = [
        "ID", "Original Title", "Chinese Title", "Regulation Number", "Jurisdiction", "Country",
        "State/Province", "Authority", "Document Type", "Legal Status", "Publication Date",
        "Entry into Force", "Application Date", "Compliance Deadline", "Topics", "Business Lines",
        "Relevance", "Chinese Summary", "Cleva Impact", "Official URL", "Source", "Source Level",
        "Last Verified", "Reviewer", "Updated At",
    ]
    ws.append(headers)
    for row in rows:
        ws.append([
            row.get("id"), row.get("original_title"), row.get("chinese_title"), row.get("regulation_number"),
            row.get("jurisdiction"), row.get("country"), row.get("state_province"), row.get("authority"),
            row.get("document_type"), row.get("legal_status"), row.get("publication_date"),
            row.get("entry_into_force_date"), row.get("application_date"), row.get("compliance_deadline"),
            _json_list(row.get("topics_json")), _json_list(row.get("business_lines_json")),
            row.get("relevance_level"), row.get("summary_cn"), row.get("cleva_impact"), row.get("official_url"),
            row.get("source_name"), row.get("source_level"), row.get("last_verified_date"), row.get("reviewer"),
            row.get("updated_at"),
        ])
    fill = PatternFill("solid", fgColor="1F4E78")
    for cell in ws[1]:
        cell.fill = fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    widths = {1: 8, 2: 42, 3: 34, 4: 22, 5: 18, 8: 26, 17: 12, 18: 50, 19: 50, 20: 55}
    for idx in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(idx)].width = widths.get(idx, 18)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def regulation_to_docx(row: Mapping[str, Any]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = row.get("chinese_title") or row.get("original_title") or "Regulation Record"
    p = doc.add_paragraph()
    run = p.add_run(str(title))
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph(str(row.get("original_title") or ""))
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    fields = [
        ("Regulation Number", row.get("regulation_number")),
        ("Jurisdiction", row.get("jurisdiction")),
        ("Country / State", " / ".join(x for x in [row.get("country"), row.get("state_province")] if x)),
        ("Authority", row.get("authority")),
        ("Document Type", row.get("document_type")),
        ("Legal Status", row.get("legal_status")),
        ("Publication Date", row.get("publication_date")),
        ("Entry into Force", row.get("entry_into_force_date")),
        ("Application Date", row.get("application_date")),
        ("Compliance Deadline", row.get("compliance_deadline")),
        ("Topics", _json_list(row.get("topics_json"))),
        ("Business Lines", _json_list(row.get("business_lines_json"))),
        ("Relevance", row.get("relevance_level")),
        ("Official URL", row.get("official_url")),
        ("Source / Level", f"{row.get('source_name') or ''} / {row.get('source_level') or ''}"),
        ("Last Verified", row.get("last_verified_date")),
        ("Reviewer", row.get("reviewer")),
    ]
    for label, value in fields:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value or "")
        cells[0].paragraphs[0].runs[0].bold = True

    doc.add_heading("法规中文摘要", level=1)
    doc.add_paragraph(str(row.get("summary_cn") or "待补充"))
    doc.add_heading("对Cleva的初步影响", level=1)
    doc.add_paragraph(str(row.get("cleva_impact") or "待人工评估"))
    doc.add_heading("审核说明", level=1)
    doc.add_paragraph(
        "本记录由系统搜索和AI初步提取生成，必须以官方原文及人工审核结论为准。"
        f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}。"
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def intelligence_to_xlsx(rows: Iterable[Mapping[str, Any]]) -> bytes:
    rows = list(rows)
    wb = Workbook()
    ws = wb.active
    ws.title = "Regulatory_Intelligence"
    headers = [
        "ID", "Original Title", "Chinese Title", "Jurisdiction", "Publication Date",
        "Topics", "Business Lines", "Relevance", "Chinese Summary", "Cleva Impact",
        "Source URL", "Source Name", "Source Type", "Source Level", "Related Official URL",
        "Last Verified", "Reviewer", "Updated At",
    ]
    ws.append(headers)
    for row in rows:
        ws.append([
            row.get("id"), row.get("original_title"), row.get("chinese_title"), row.get("jurisdiction"),
            row.get("publication_date"), _json_list(row.get("topics_json")),
            _json_list(row.get("business_lines_json")), row.get("relevance_level"),
            row.get("summary_cn"), row.get("cleva_impact"), row.get("source_url"),
            row.get("source_name"), row.get("source_type"), row.get("source_level"),
            row.get("related_official_url"), row.get("last_verified_date"), row.get("reviewer"),
            row.get("updated_at"),
        ])
    fill = PatternFill("solid", fgColor="548235")
    for cell in ws[1]:
        cell.fill = fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    widths = {1: 8, 2: 42, 3: 34, 4: 18, 6: 26, 9: 50, 10: 50, 11: 55, 15: 55}
    for idx in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(idx)].width = widths.get(idx, 18)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()
